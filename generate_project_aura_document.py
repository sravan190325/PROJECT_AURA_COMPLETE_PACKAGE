from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_FILE = "Project_Aura_Detailed_Explanation_Document.docx"


BLUE = "053057"
TEAL = "00A7A7"
LIGHT_BLUE = "EAF4FB"
LIGHT_GRAY = "F4F6F8"
DARK_GRAY = "333333"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK_GRAY):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr_cells[i], BLUE)
        if widths:
            hdr_cells[i].width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]

    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    return paragraph


def add_page_break(doc):
    doc.add_page_break()


def create_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK_GRAY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT AURA")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Detailed Architecture, Page-by-Page Explanation, AI Glossary, and Deployment Notes")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    doc.add_paragraph()
    add_body(
        doc,
        "Prepared for explaining the Project Aura complete package. This document is based on the local repository structure, "
        "README, architecture documentation, route files, service files, configuration, and deployment artifacts present in the package.",
    )
    add_body(
        doc,
        "Important Harness note: the local package does not show Harness pipeline files or execution logs. The Harness section in this "
        "document explains the Harness-ready CI/CD design and how Harness would be used in production. It should not be presented as proof "
        "that Harness was actually executed unless separate Harness evidence exists.",
    )

    add_table(
        doc,
        ["Document Area", "What It Explains"],
        [
            ["Business overview", "What Project Aura does and why it is useful for PMO/project planning teams."],
            ["Architecture", "How frontend pages, Flask routes, services, AI analysis, database, and workbook generation work together."],
            ["Each page/screen", "Purpose, user actions, backend APIs, data shown, and expected outcomes for every main UI page."],
            ["AI terminology", "Plain-English definitions of AI, LLM, Claude, prompt, token, context, RAG, embeddings, validation, and related terms."],
            ["Technical terminology", "Definitions of Flask, route, API, JSON, SQLite, session, environment variables, Docker, and CI/CD."],
            ["Harness", "A truthful Harness-ready deployment design and presentation wording."],
        ],
        widths=[Inches(2.0), Inches(4.8)],
    )

    add_page_break(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "Project Aura is an AI-powered project planning assistant. It takes project input documents such as SOWs, requirement documents, "
        "PDFs, DOCX files, or PPTX decks, extracts their text, analyzes the content with Claude AI or demo/mock logic, stores structured "
        "project information in SQLite, and generates professional project management workbooks.",
    )
    add_body(
        doc,
        "The main business value is faster project plan creation. Instead of manually reading long SOWs and building PMO templates from scratch, "
        "the system produces structured outputs such as project details, deliverables, staffing plans, risks, milestones, RACI, budget trackers, "
        "dashboards, and workbook tabs that can be used by delivery managers and stakeholders.",
    )
    add_bullet(doc, "Primary users: project managers, PMO analysts, delivery leads, solution architects, and consulting teams.")
    add_bullet(doc, "Primary inputs: PDF, DOCX, and PPTX project documents.")
    add_bullet(doc, "Primary processing: file validation, text extraction, AI analysis, clarification, database save, workbook generation.")
    add_bullet(doc, "Primary outputs: project summary pages and Excel workbooks saved in the workbooks folder.")
    add_bullet(doc, "Main stack: Flask, Python, SQLite, Bootstrap/Blend templates, JavaScript, Claude API, openpyxl, pdfplumber, python-docx, python-pptx.")

    add_heading(doc, "2. What The Application Does", 1)
    add_number(doc, "The user opens the Project Aura home page.")
    add_number(doc, "The user uploads one or more project documents.")
    add_number(doc, "The backend validates the file extension and saves the upload securely.")
    add_number(doc, "The document processor routes the file to the right parser based on type: PDF, DOCX, or PPTX.")
    add_number(doc, "Extracted text is saved into a companion text file to avoid overloading the browser session cookie.")
    add_number(doc, "The user reviews the extracted content on the results page.")
    add_number(doc, "The project documents are analyzed with Claude AI, demo data, or mock mode depending on configuration.")
    add_number(doc, "The user confirms missing planning fields such as start date, duration, team size, and delivery model.")
    add_number(doc, "The project is saved into SQLite tables for projects, documents, deliverables, team members, and risks.")
    add_number(doc, "The workbook generator creates a project planning workbook and makes it available for download.")

    add_heading(doc, "3. High-Level Artifact Diagram", 1)
    add_body(doc, "The following Mermaid diagram can be copied into any Mermaid renderer to visualize the project artifacts.")
    add_code(
        doc,
        """flowchart TD
    User["User / PMO Analyst"]

    subgraph Inputs["Input Artifacts"]
        PDF["PDF"]
        DOCX["DOCX"]
        PPTX["PPTX"]
        SOW["SOW / Requirements"]
    end

    subgraph WebApp["Project Aura Flask Application"]
        App["app.py"]
        Routes["routes/"]
        Services["services/"]
        Templates["templates/"]
        Static["static/"]
        Config["config.py + .env"]
    end

    subgraph Processing["Processing Artifacts"]
        Uploads["uploads/"]
        Temp["temp/"]
        Text["Extracted Text"]
        AI["Claude / Mock AI Analysis"]
        Structured["Structured Project Data"]
    end

    subgraph Data["SQLite Database"]
        DB["project_aura.db"]
        Projects["projects"]
        Documents["documents"]
        Deliverables["deliverables"]
        Team["team_members"]
        Risks["risks"]
    end

    subgraph Outputs["Output Artifacts"]
        Summary["Project Summary Page"]
        Workbook["Excel Project Workbook"]
        Platform["Optional Platform Delivery: Excel / Smartsheet / Jira"]
    end

    User --> SOW
    SOW --> PDF
    SOW --> DOCX
    SOW --> PPTX
    PDF --> Uploads
    DOCX --> Uploads
    PPTX --> Uploads
    Uploads --> Text
    Text --> AI
    AI --> Structured
    Structured --> DB
    DB --> Projects
    DB --> Documents
    DB --> Deliverables
    DB --> Team
    DB --> Risks
    Structured --> Summary
    Structured --> Workbook
    Structured --> Platform
    Workbook --> User""",
    )

    add_page_break(doc)

    add_heading(doc, "4. Repository Structure", 1)
    add_table(
        doc,
        ["Folder/File", "Purpose"],
        [
            ["app.py", "Main Flask entry point. Loads config, creates the Flask app, registers route blueprints, creates folders, and starts the server."],
            ["config.py", "Central configuration for secret key, debug mode, session behavior, upload folder, temp folder, allowed file types, and extraction limits."],
            ["routes/", "Defines HTTP routes and APIs for uploads, project analysis, workbook generation, and platform delivery."],
            ["services/", "Contains business logic for document parsing, Claude analysis, project detection, database operations, workbook creation, Gantt generation, and resource matching."],
            ["templates/", "HTML pages rendered by Flask. Includes Blend-styled and older legacy templates."],
            ["static/css/", "CSS files for styling. blend.css is the newer Blend-style UI file and style.css is the original styling."],
            ["static/js/upload.js", "Client-side upload behavior, drag-and-drop handling, upload progress, and front-end API calls."],
            ["uploads/", "Temporary uploaded source documents."],
            ["temp/", "Temporary working files used during processing."],
            ["workbooks/", "Generated Excel workbook output files."],
            ["data/talent_pool.csv", "Sample talent pool used by the resource recommender."],
            ["project_aura.db", "SQLite database used by the local application."],
            ["Dockerfile, docker-compose.yml, Procfile, runtime.txt", "Deployment artifacts for container/cloud deployment."],
            ["requirements.txt", "Python dependency list."],
        ],
        widths=[Inches(2.3), Inches(4.7)],
    )

    add_heading(doc, "5. Technology Stack", 1)
    add_table(
        doc,
        ["Technology", "How Project Aura Uses It"],
        [
            ["Flask 2.3.3", "Python web framework used to serve pages and APIs."],
            ["Werkzeug", "Provides secure filename handling and WSGI utilities used by Flask."],
            ["pdfplumber", "Extracts text and metadata from PDF files."],
            ["python-docx", "Extracts text from DOCX files and can also create Word documents."],
            ["python-pptx", "Extracts text from PowerPoint PPTX slides."],
            ["openpyxl", "Creates and formats Excel workbook outputs."],
            ["anthropic", "Official Python client used to call Claude AI."],
            ["SQLite", "Local file-based database for projects, documents, deliverables, team members, and risks."],
            ["Bootstrap / Blend CSS", "Provides responsive UI styling and layout."],
            ["Vanilla JavaScript", "Handles browser upload interactions and calls Flask APIs."],
            ["Docker / Procfile / Gunicorn", "Deployment packaging and production server options."],
        ],
        widths=[Inches(1.9), Inches(5.1)],
    )

    add_heading(doc, "6. End-to-End User Journey", 1)
    add_number(doc, "Home page: user lands on the upload interface.")
    add_number(doc, "Upload: user selects PDF, DOCX, or PPTX files.")
    add_number(doc, "Extraction: backend extracts text and metadata.")
    add_number(doc, "Results page: user reviews extracted text and document metadata.")
    add_number(doc, "AI analysis: user triggers analysis; Claude returns project type, scope, deliverables, team needs, confidence, and missing fields.")
    add_number(doc, "Clarification page: user supplies missing planning information.")
    add_number(doc, "Project creation: backend saves the final project into SQLite.")
    add_number(doc, "Summary page: user sees a structured project overview, KPIs, roles, recommendations, risks, and next actions.")
    add_number(doc, "Workbook generation: backend creates the Excel project workbook.")
    add_number(doc, "Download or platform delivery: user downloads Excel or optionally creates project artifacts in supported platforms.")

    add_page_break(doc)

    add_heading(doc, "7. Page-by-Page Explanation", 1)
    add_body(
        doc,
        "This section explains every major page/template in the package. The Blend templates are the newer styled versions used by the current app routes where applicable. "
        "The non-Blend templates remain as legacy equivalents or fallback pages.",
    )

    add_table(
        doc,
        ["Page/Template", "Route", "Purpose", "What The User Does"],
        [
            ["base_blend.html", "Layout template", "Shared layout used by Blend-styled pages. It defines the common page frame, styling includes, navigation, and shared UI structure.", "The user does not open this directly; other pages extend or reuse it."],
            ["base.html", "Layout template", "Older shared base template for legacy pages.", "Not usually opened directly."],
            ["index_blend.html", "/", "Main home and upload page. It introduces Project Aura and provides the document upload interface.", "User drags or selects PDF, DOCX, or PPTX files and starts upload."],
            ["index.html", "Legacy home", "Original home/upload page before the Blend-styled version.", "Useful as fallback or reference."],
            ["results_blend.html", "/api/results", "Shows extracted document content after upload.", "User reviews preview, metadata, and clicks the analysis action."],
            ["results.html", "Legacy results", "Older extracted-content results page.", "Fallback/reference page."],
            ["clarification_blend.html", "/api/project/clarify", "Collects missing project planning fields after AI analysis.", "User enters start date, duration, team size, and delivery model."],
            ["clarification.html", "Legacy clarify", "Older clarification form.", "Fallback/reference page."],
            ["project_summary_blend.html", "/api/project/<id>/summary", "Displays the created project summary and project planning outputs.", "User reviews project details, metrics, resource recommendations, risks, and workbook actions."],
            ["project_summary.html", "Legacy summary", "Older project summary page.", "Fallback/reference page."],
            ["workbook_generator.html", "Workbook UI", "Page for workbook generation and preview actions.", "User can request workbook generation, preview sheet list, and download output."],
            ["platform_delivery_home.html", "Defined as /platform-delivery but delivery_bp is not registered in app.py", "Home page for multi-platform delivery workflow.", "Would let user choose a project for platform delivery if the blueprint is registered."],
            ["platform_delivery_selection.html", "/api/platform/selection/<id>", "Platform selection page for Excel, Smartsheet, and Jira delivery.", "User selects a target platform, tests credentials where required, and creates the project in that platform."],
        ],
        widths=[Inches(1.7), Inches(1.5), Inches(2.4), Inches(1.9)],
    )

    add_heading(doc, "7.1 Home / Upload Page", 2)
    add_body(
        doc,
        "The home page is served by app.py at the root route '/'. It renders index_blend.html. This is the first screen the user sees. "
        "The page is designed for quick document intake and supports drag-and-drop or file picker uploads.",
    )
    add_bullet(doc, "Inputs accepted: PDF, DOCX, and PPTX.")
    add_bullet(doc, "Backend API called: POST /api/upload.")
    add_bullet(doc, "Validation: file extension is checked against Config.ALLOWED_EXTENSIONS.")
    add_bullet(doc, "Storage: files are saved in uploads/ with secure and unique filenames.")
    add_bullet(doc, "Output: upload response includes processed count, failed count, document metadata, preview text, page count, and table presence.")

    add_heading(doc, "7.2 Results Page", 2)
    add_body(
        doc,
        "The results page is rendered by /api/results after upload. It loads processed documents from the Flask session and reads extracted text from companion text files. "
        "This page proves that the document extraction step worked before the AI step runs.",
    )
    add_bullet(doc, "Shows uploaded file names, file types, metadata, extracted text preview, and processing status.")
    add_bullet(doc, "Allows the user to proceed to AI analysis.")
    add_bullet(doc, "Keeps the workflow transparent because users can verify extracted content before generating project plans.")

    add_heading(doc, "7.3 Clarification Page", 2)
    add_body(
        doc,
        "The clarification page is served by GET /api/project/clarify. It appears after AI analysis and asks the user to provide fields that are usually missing or ambiguous in an SOW. "
        "The route redirects to home if no analysis exists in session.",
    )
    add_bullet(doc, "Fields: start date, project duration in weeks, team size, and delivery model.")
    add_bullet(doc, "Backend API on submit: POST /api/project/clarify.")
    add_bullet(doc, "Purpose: merge AI-extracted information with human-confirmed planning assumptions.")
    add_bullet(doc, "Output: creates a project record, stores documents, deliverables, staffing, and risks, then returns the new project ID.")

    add_heading(doc, "7.4 Project Summary Page", 2)
    add_body(
        doc,
        "The project summary page is rendered by GET /api/project/<project_id>/summary. It reads complete project data from SQLite and formats it for the user. "
        "It also calculates end date from start date and duration where possible.",
    )
    add_bullet(doc, "Shows project name, project type, client, scope, timeline, duration, team size, and delivery model.")
    add_bullet(doc, "Shows counts for deliverables, team members, and risks.")
    add_bullet(doc, "Displays resource recommendations from data/talent_pool.csv using role and skill matching.")
    add_bullet(doc, "Acts as the main project workspace after a project is created.")

    add_heading(doc, "7.5 Workbook Generator Page", 2)
    add_body(
        doc,
        "The workbook generator page supports the final export workflow. The API routes behind it are POST /api/workbook/generate/<project_id>, "
        "GET /api/workbook/preview/<project_id>, and GET /api/workbook/download/<project_id>.",
    )
    add_bullet(doc, "Preview tells the user what sheets will be generated.")
    add_bullet(doc, "Generate creates the workbook in the workbooks/ folder.")
    add_bullet(doc, "Download streams the Excel file back to the browser.")
    add_bullet(doc, "The default generator uses the sample/improved workbook format unless a different generator query parameter is passed.")

    add_heading(doc, "7.6 Platform Delivery Pages", 2)
    add_body(
        doc,
        "The package includes platform delivery screens for creating project outputs in Excel, Smartsheet, or Jira. The API blueprint is registered, so /api/platform/selection/<project_id>, "
        "/api/platform/test-connection, and /api/project/create-platform are available. The separate delivery_bp home route is defined in code but is not registered in app.py in the current package.",
    )
    add_bullet(doc, "Excel platform: creates an Excel artifact using the platform creator.")
    add_bullet(doc, "Smartsheet platform: tests a Smartsheet token and creates a sheet/project if credentials are valid.")
    add_bullet(doc, "Jira platform: tests Jira connection details and creates a Scrum-style project output if credentials are valid.")
    add_bullet(doc, "Platform mappings are logged, and the code includes a TODO for permanent mapping storage.")

    add_page_break(doc)

    add_heading(doc, "8. Backend API Explanation", 1)
    add_table(
        doc,
        ["API", "Method", "Purpose", "Input", "Output"],
        [
            ["/health", "GET", "Checks whether the Flask server is alive.", "None", "JSON status healthy."],
            ["/api/upload", "POST", "Uploads and processes one or more documents.", "Multipart files field named files.", "JSON results with preview, metadata, errors, and total document count."],
            ["/api/documents", "GET", "Returns processed documents for the current session.", "Current browser session.", "JSON document list including extracted text loaded from text files."],
            ["/api/clear", "POST", "Clears uploaded documents from the session and removes temp files.", "Current browser session.", "JSON success message."],
            ["/api/results", "GET", "Renders the extraction results page.", "Current browser session.", "HTML results page."],
            ["/api/project/analyze", "POST", "Runs AI/demo analysis on uploaded documents.", "Processed documents in session.", "JSON analysis, validation, risks, resource recommendations, and project config."],
            ["/api/project/clarify", "GET", "Shows the clarification form.", "Analysis in session.", "HTML clarification page."],
            ["/api/project/clarify", "POST", "Saves clarified project data.", "start_date, duration, team_size, delivery_model.", "JSON project ID and success response."],
            ["/api/project/list", "GET", "Lists all projects from SQLite.", "None", "JSON project list."],
            ["/api/project/<id>", "GET", "Gets full project data by ID.", "Project ID.", "JSON project summary or 404."],
            ["/api/project/<id>/summary", "GET", "Renders project summary page.", "Project ID.", "HTML summary page."],
            ["/api/workbook/generate/<id>", "POST", "Generates Excel workbook.", "Project ID and optional generator query parameter.", "JSON filename and download URL."],
            ["/api/workbook/download/<id>", "GET", "Downloads or regenerates workbook.", "Project ID and optional generator query parameter.", "XLSX file stream."],
            ["/api/workbook/preview/<id>", "GET", "Shows workbook sheet preview.", "Project ID.", "JSON sheet list and stats."],
            ["/api/platform/selection/<id>", "GET", "Shows platform selection page.", "Project ID.", "HTML platform selection page."],
            ["/api/platform/test-connection", "POST", "Tests Smartsheet or Jira credentials.", "Platform and credentials JSON.", "JSON connection result."],
            ["/api/project/create-platform", "POST", "Creates project artifact in Excel, Smartsheet, or Jira.", "Project ID, platform, credentials.", "JSON result and created platform reference."],
        ],
        widths=[Inches(1.6), Inches(0.7), Inches(2.1), Inches(1.6), Inches(1.6)],
    )

    add_heading(doc, "9. Service Layer Explanation", 1)
    add_table(
        doc,
        ["Service", "Responsibility"],
        [
            ["DocumentProcessor", "Central router for file parsing. It checks file type and sends PDFs to PDFParser, DOCX files to DOCXParser, and PPTX files to PPTXParser."],
            ["PDFParser", "Extracts text, metadata, pages, and optionally table content from PDF files using pdfplumber."],
            ["DOCXParser", "Extracts paragraph and document content from Word files using python-docx."],
            ["PPTXParser", "Extracts text from PowerPoint slides using python-pptx."],
            ["ClaudeService", "Builds prompts, calls Claude, parses JSON responses, and supports mock mode when real API calls are not desired."],
            ["ProjectDetector", "Validates project information, detects project type, estimates staffing, identifies risks, and provides phase templates."],
            ["DatabaseService", "Creates SQLite tables and performs CRUD operations for projects, documents, deliverables, team members, and risks."],
            ["ResourceRecommender", "Matches required project roles against the talent pool CSV using role, skills, availability, experience, and domain fit."],
            ["ProjectPlanEngine", "Creates project plan structures, phases, dates, dependencies, and scheduling logic for workbook generation."],
            ["ExcelFormatter", "Applies professional Excel styling such as fonts, colors, borders, column widths, headers, and formatting rules."],
            ["Workbook generators", "Create Excel workbooks in standard, enhanced, PMO, optimized, and sample formats."],
            ["Gantt generators", "Generate Gantt-style schedule views and project timeline artifacts."],
            ["Platform creators", "Create project outputs for Excel, Smartsheet, and Jira in the multi-platform delivery workflow."],
        ],
        widths=[Inches(2.0), Inches(5.0)],
    )

    add_heading(doc, "10. Data Flow In Detail", 1)
    add_number(doc, "Browser sends files to /api/upload.")
    add_number(doc, "upload_routes.py checks that the request includes files.")
    add_number(doc, "Each file name is sanitized using secure_filename to reduce path and naming risks.")
    add_number(doc, "A UUID is prefixed to the file name to prevent collisions between uploads.")
    add_number(doc, "The file is saved to Config.UPLOAD_FOLDER.")
    add_number(doc, "DocumentProcessor validates the file extension and selects the parser.")
    add_number(doc, "The parser extracts text and metadata from PDF, DOCX, or PPTX.")
    add_number(doc, "Extracted text is saved into a .txt companion file beside the uploaded file.")
    add_number(doc, "Lightweight metadata and text path are stored in the Flask session.")
    add_number(doc, "The results page loads the extracted text from disk and displays it.")
    add_number(doc, "ClaudeService prepares the text and builds a structured prompt.")
    add_number(doc, "Claude returns structured JSON including project type, scope, deliverables, team requirements, confidence, and missing fields.")
    add_number(doc, "ProjectDetector validates the AI output and adds project type configuration, staffing, phases, and risks.")
    add_number(doc, "The clarification form collects missing human planning inputs.")
    add_number(doc, "DatabaseService saves the final project and related records.")
    add_number(doc, "Workbook routes retrieve saved project data and generate an Excel workbook.")

    add_page_break(doc)

    add_heading(doc, "11. Database Design", 1)
    add_body(
        doc,
        "Project Aura uses SQLite locally through project_aura.db. SQLite is a file-based relational database, which makes it simple for a local demo or small deployment. "
        "The database is initialized automatically when DatabaseService starts.",
    )
    add_table(
        doc,
        ["Table", "Main Columns", "Purpose"],
        [
            ["projects", "id, project_name, project_type, client_name, scope, start_date, duration_weeks, team_size, delivery_model, status, extracted_data, created_at, updated_at", "Stores the master record for each generated project."],
            ["documents", "id, project_id, filename, file_type, extracted_text, file_size, uploaded_at", "Stores uploaded document records and extracted text linked to a project."],
            ["deliverables", "id, project_id, deliverable_name, description, status", "Stores deliverables extracted from AI analysis or project configuration."],
            ["team_members", "id, project_id, role, count, resource_allocated", "Stores team role breakdown, such as PM, developer, QA, data engineer, and architect."],
            ["risks", "id, project_id, risk_description, severity, mitigation, status", "Stores identified risks and suggested mitigation steps."],
        ],
        widths=[Inches(1.3), Inches(3.6), Inches(2.1)],
    )
    add_body(
        doc,
        "Relationship explanation: one project can have many documents, many deliverables, many team member role records, and many risks. "
        "The project_id foreign key links child records back to the main project.",
    )

    add_heading(doc, "12. AI Processing Explanation", 1)
    add_body(
        doc,
        "The AI flow is handled mainly by ClaudeService and ProjectDetector. ClaudeService prepares the uploaded document text, builds prompts, calls Claude through the Anthropic API, "
        "and parses the JSON response. ProjectDetector validates and enriches the result with business rules.",
    )
    add_bullet(doc, "The model configured in code is claude-3-5-sonnet-20241022.")
    add_bullet(doc, "The prompt asks for project_type, client_name, scope, deliverables, team_requirements, confidence, missing_fields, and key_highlights.")
    add_bullet(doc, "The service limits the amount of content sent to Claude so very large documents do not overload the prompt.")
    add_bullet(doc, "Mock mode exists through USE_MOCK_CLAUDE. This allows local testing without calling the real Claude API.")
    add_bullet(doc, "Demo mode exists through demo_mode.py. It can return sample project data for demonstrations.")
    add_bullet(doc, "The AI output is expected as JSON so the backend can parse and store it.")

    add_heading(doc, "13. AI Terminology Glossary", 1)
    add_table(
        doc,
        ["Term", "Meaning In Simple Words", "How Project Aura Uses It"],
        [
            ["AI", "Artificial Intelligence. Software that performs tasks that normally require human reasoning or pattern recognition.", "Used to read project documents and infer structured planning information."],
            ["LLM", "Large Language Model. A model trained on large amounts of text that can understand and generate language.", "Claude is the LLM used to analyze uploaded project documents."],
            ["Claude", "Anthropic's LLM family.", "Called through the anthropic Python package for document analysis."],
            ["Prompt", "The instruction sent to an AI model.", "ClaudeService builds prompts asking Claude to return project details in JSON."],
            ["Token", "A small chunk of text processed by an AI model.", "Long documents must be shortened because model calls have token limits."],
            ["Context", "The information provided to the model for a specific request.", "Uploaded document text and user clarifications become the context for analysis."],
            ["Context window", "Maximum amount of text the model can consider in one request.", "The service limits content length before sending text to Claude."],
            ["JSON", "A structured text format for data exchange.", "Claude is instructed to return JSON so Python can parse project fields reliably."],
            ["Classification", "Choosing a category for something.", "The app classifies project type, such as Data Engineering, GenAI, Cloud Migration, or Application Development."],
            ["Information extraction", "Finding specific fields inside unstructured text.", "The app extracts client name, scope, deliverables, team needs, and missing fields."],
            ["Confidence score", "A numeric estimate of how sure the analysis is.", "Returned by Claude or mock logic to indicate certainty of project type detection."],
            ["NLP", "Natural Language Processing. AI methods for working with human language.", "Used conceptually when the system understands SOW language."],
            ["RAG", "Retrieval-Augmented Generation. A pattern where the system retrieves relevant data before asking the model to generate an answer.", "The architecture documentation describes RAG as a future/enterprise pattern, but the local code does not show a live vector database implementation."],
            ["Embedding", "A numeric vector that represents the meaning of text.", "Described in architecture docs for similarity search; not implemented as a local vector store in the scanned package."],
            ["Vector database", "A database optimized for searching embeddings by similarity.", "Referenced in architecture docs as a scalable production enhancement."],
            ["Semantic search", "Search based on meaning rather than exact keywords.", "Would help find similar past projects in a RAG design."],
            ["Hallucination", "When an AI model generates incorrect or unsupported information.", "Reduced by asking for JSON, validating fields, and requiring user clarification."],
            ["Validation", "Checking whether output is complete and reasonable.", "ProjectDetector checks extracted project information and identifies missing fields."],
            ["Mock mode", "A test mode that returns generated sample analysis instead of calling the real AI API.", "Useful when API quota is unavailable or a demo must run offline."],
        ],
        widths=[Inches(1.3), Inches(2.7), Inches(3.0)],
    )

    add_page_break(doc)

    add_heading(doc, "14. Technical Terminology Glossary", 1)
    add_table(
        doc,
        ["Term", "Explanation", "Example In Project Aura"],
        [
            ["Flask", "A lightweight Python web framework.", "app.py creates the Flask app."],
            ["Route", "A URL path handled by backend code.", "/api/upload handles file uploads."],
            ["Blueprint", "A Flask feature for grouping routes by feature area.", "upload_bp, project_bp, workbook_bp, platform_bp."],
            ["API", "Application Programming Interface. A structured way for software parts to communicate.", "Browser JavaScript calls /api/project/analyze."],
            ["GET", "HTTP method usually used to read data or render pages.", "GET /api/project/list returns projects."],
            ["POST", "HTTP method usually used to submit or create data.", "POST /api/upload submits files."],
            ["Session", "Server-side/browser-associated state for one user session.", "Stores processed_documents and project_analysis between pages."],
            ["Cookie", "Small browser storage used to identify a session.", "Session references rely on cookies, but large extracted text is stored in files to avoid cookie size overflow."],
            ["Environment variable", "Configuration value stored outside source code.", "ANTHROPIC_API_KEY and SECRET_KEY are loaded from .env."],
            [".env", "Local environment configuration file.", "Stores the Claude API key and Flask environment settings."],
            ["SQLite", "Small file-based relational database.", "project_aura.db stores project records."],
            ["CRUD", "Create, Read, Update, Delete database operations.", "DatabaseService creates and reads project records."],
            ["Template", "HTML file rendered by Flask with dynamic data.", "project_summary_blend.html receives project data."],
            ["Static asset", "CSS, JS, image, or other file served directly to the browser.", "blend.css and upload.js."],
            ["MIME type", "A content type that identifies file format.", "PDF files use application/pdf."],
            ["Secure filename", "A cleaned file name safe for storage.", "Werkzeug secure_filename is used before saving uploads."],
            ["UUID", "A unique identifier.", "Used to prefix uploaded file names and avoid collisions."],
            ["Dockerfile", "Instructions to build a container image.", "Used for production packaging."],
            ["docker-compose", "Local multi-container orchestration file.", "Can run app and supporting services together."],
            ["Gunicorn", "Production WSGI server for Python web apps.", "Listed in requirements for deployment."],
            ["Procfile", "Cloud platform startup command file.", "Useful on platforms such as Render or Heroku-style deploys."],
            ["CI/CD", "Continuous Integration and Continuous Delivery/Deployment.", "Could be implemented with Harness for automated build, test, and deploy."],
        ],
        widths=[Inches(1.4), Inches(2.8), Inches(2.8)],
    )

    add_heading(doc, "15. Workbook Output Explanation", 1)
    add_body(
        doc,
        "The workbook output is the main final artifact. The workbook routes can generate different workbook formats. The current default preview lists the approved sample/improved workbook format.",
    )
    add_table(
        doc,
        ["Workbook Sheet", "What It Means"],
        [
            ["00_Home", "Entry sheet or navigation page for the workbook."],
            ["Executive_Dashboard", "High-level summary for leadership, usually including health, key metrics, risks, and status."],
            ["Detailed_Task_Plan", "Detailed plan of tasks, owners, dates, priorities, and dependencies."],
            ["Budget_Tracker", "Tracks planned and actual budget values."],
            ["Milestone_Tracker", "Tracks major milestone dates and completion status."],
            ["Resource_Plan", "Shows required roles, staffing counts, allocations, and resource needs."],
            ["RAID_Register", "Captures Risks, Assumptions, Issues, and Dependencies."],
            ["RACI_Matrix", "Shows who is Responsible, Accountable, Consulted, and Informed for each activity."],
            ["Change_Log", "Tracks approved changes to scope, timeline, budget, or requirements."],
        ],
        widths=[Inches(2.0), Inches(5.0)],
    )
    add_body(
        doc,
        "Some older documentation mentions a 14-sheet workbook or a 12-sheet PMO workbook. The code currently previews a 9-sheet sample/improved format by default, while other generator classes still exist for legacy or alternate workbook styles.",
    )

    add_heading(doc, "16. Security And Configuration", 1)
    add_bullet(doc, "API keys are expected in environment variables instead of hardcoded source code.")
    add_bullet(doc, "SECRET_KEY is configurable and should be changed for production.")
    add_bullet(doc, "SESSION_COOKIE_HTTPONLY is enabled to reduce client-side script access to the session cookie.")
    add_bullet(doc, "SESSION_COOKIE_SECURE is enabled in ProductionConfig and should be used with HTTPS.")
    add_bullet(doc, "Allowed uploads are restricted to PDF, DOCX, and PPTX.")
    add_bullet(doc, "File names are sanitized with Werkzeug secure_filename.")
    add_bullet(doc, "Generated and uploaded files should be managed carefully in production with retention policies.")

    add_page_break(doc)

    add_heading(doc, "17. Harness CI/CD Integration Page", 1)
    add_body(
        doc,
        "Harness is a CI/CD and software delivery platform. It can automate building, testing, scanning, approving, deploying, monitoring, and rolling back applications. "
        "For Project Aura, Harness would sit outside the Flask application and manage the deployment workflow from source code to runtime environment.",
    )
    add_body(
        doc,
        "Accuracy note for presentation: this local package includes deployment artifacts such as Dockerfile, docker-compose.yml, Procfile, runtime.txt, and requirements.txt, "
        "but it does not include Harness pipeline YAML or execution logs. The correct wording is that the project is Harness-ready or that Harness can be used as the CI/CD layer.",
    )
    add_table(
        doc,
        ["Harness Stage", "What It Would Do For Project Aura"],
        [
            ["Source trigger", "Start the pipeline when code is pushed to GitHub, GitLab, or another repository."],
            ["Dependency install", "Install Python dependencies from requirements.txt."],
            ["Unit and smoke tests", "Run Flask health checks, parser tests, and workbook generation checks."],
            ["Secret scan", "Check that .env secrets such as ANTHROPIC_API_KEY are not committed."],
            ["Build container", "Build a Docker image using the Dockerfile."],
            ["Security scan", "Scan the image for vulnerable packages and known CVEs."],
            ["Deploy to dev", "Deploy to a development environment for validation."],
            ["Approval gate", "Require PM/lead approval before production deployment."],
            ["Deploy to production", "Deploy the approved container to Render, AWS ECS, Azure App Service, Kubernetes, or another target."],
            ["Post-deploy verification", "Call /health and run a sample upload/analyze workflow."],
            ["Rollback", "Return to the previous successful version if production validation fails."],
        ],
        widths=[Inches(2.0), Inches(5.0)],
    )

    add_heading(doc, "17.1 Harness Pipeline Design", 2)
    add_number(doc, "Developer commits Project Aura code to the repository.")
    add_number(doc, "Harness detects the commit and starts the pipeline.")
    add_number(doc, "Harness installs dependencies and runs automated checks.")
    add_number(doc, "Harness builds a Docker image using the project Dockerfile.")
    add_number(doc, "Harness stores secrets securely, including ANTHROPIC_API_KEY and SECRET_KEY.")
    add_number(doc, "Harness deploys the image into a dev or staging environment.")
    add_number(doc, "A manual approval gate protects production deployment.")
    add_number(doc, "Harness deploys production and verifies /health.")
    add_number(doc, "Harness keeps deployment history and supports rollback.")

    add_heading(doc, "17.2 Presentation-Safe Harness Wording", 2)
    add_body(
        doc,
        "Recommended wording: Project Aura was designed to be deployment-ready with container and cloud startup artifacts. Harness can be used as the CI/CD orchestration layer to automate build, test, approval, deployment, verification, and rollback."
    )
    add_body(
        doc,
        "Avoid saying: We executed the production deployment through Harness. That statement would require actual Harness pipeline evidence."
    )
    add_body(
        doc,
        "Short slide wording: Harness-ready CI/CD design - automated source trigger, dependency install, test, Docker build, secret management, approval gate, environment deployment, health check, and rollback."
    )

    add_heading(doc, "17.3 Example Harness-Ready Flow", 2)
    add_code(
        doc,
        """Git push
  -> Harness pipeline trigger
  -> Install dependencies from requirements.txt
  -> Run tests and smoke checks
  -> Build Docker image from Dockerfile
  -> Inject secrets from Harness secret manager
  -> Deploy to dev/stage
  -> Manual approval
  -> Deploy to production
  -> Verify /health
  -> Roll back if verification fails""",
    )

    add_page_break(doc)

    add_heading(doc, "18. Deployment Explanation", 1)
    add_body(
        doc,
        "Project Aura can run locally or be packaged for deployment. Locally, the developer creates a Python virtual environment, installs requirements, sets .env values, and starts app.py. "
        "For production, the package contains Dockerfile, docker-compose.yml, Procfile, runtime.txt, and Gunicorn dependency support.",
    )
    add_table(
        doc,
        ["Artifact", "Role In Deployment"],
        [
            ["requirements.txt", "Lists Python packages to install."],
            [".env.example", "Shows required configuration keys without exposing real secrets."],
            ["Dockerfile", "Defines how to build a container image for the app."],
            ["docker-compose.yml", "Supports local container orchestration."],
            ["Procfile", "Defines a process start command for cloud platforms that use Procfile conventions."],
            ["runtime.txt", "Specifies Python runtime version expectations."],
            ["wsgi.py", "Production WSGI entry point for Gunicorn or similar servers."],
        ],
        widths=[Inches(1.8), Inches(5.2)],
    )

    add_heading(doc, "19. Error Handling And Reliability", 1)
    add_bullet(doc, "413 handler returns a clear message when uploaded files are too large.")
    add_bullet(doc, "400 handler returns a clear message for invalid requests.")
    add_bullet(doc, "500 handler hides internal implementation details behind a generic error response.")
    add_bullet(doc, "Upload route tracks per-file success and failure, so one bad file does not necessarily stop the entire upload batch.")
    add_bullet(doc, "Temporary files are cleaned when the user clears the session.")
    add_bullet(doc, "Demo/mock modes allow the system to keep working for demos even when Claude API access is unavailable.")

    add_heading(doc, "20. Known Implementation Notes", 1)
    add_bullet(doc, "The current app.py registers upload, project, workbook, platform, and project-delivery blueprints.")
    add_bullet(doc, "The delivery_bp routes for /platform-delivery are defined but not registered in app.py in the scanned package.")
    add_bullet(doc, "Architecture documentation mentions PostgreSQL, RAG, vector databases, and multi-agent architecture as enterprise architecture concepts. The local implementation primarily uses SQLite and a direct Claude analysis service.")
    add_bullet(doc, "Workbook documentation references several workbook sheet counts. The active preview route currently lists a 9-sheet sample/improved format by default.")
    add_bullet(doc, "For production, secrets, file retention, user authentication, authorization, and persistent object storage should be hardened.")

    add_page_break(doc)

    add_heading(doc, "21. Suggested Demo Script", 1)
    add_number(doc, "Open the home page and explain that Project Aura converts project documents into PMO-ready plans.")
    add_number(doc, "Upload a PDF, DOCX, or PPTX SOW.")
    add_number(doc, "Show the results page and explain that the backend extracted text before using AI.")
    add_number(doc, "Click analysis and explain that Claude classifies the project type and extracts scope, deliverables, team requirements, and missing fields.")
    add_number(doc, "Fill clarification fields: start date, duration, team size, and delivery model.")
    add_number(doc, "Show the project summary and explain database-backed records for project, documents, deliverables, team, and risks.")
    add_number(doc, "Generate the workbook and explain the sheets: dashboard, detailed task plan, budget, milestones, resources, RAID, RACI, and change log.")
    add_number(doc, "Mention Harness as the recommended CI/CD layer for automated build, test, approval, deployment, verification, and rollback.")

    add_heading(doc, "22. Viva / Interview Questions And Answers", 1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["What problem does Project Aura solve?", "It reduces manual effort in converting SOWs and project documents into structured project plans and Excel workbooks."],
            ["Why use AI?", "AI helps read unstructured project text and convert it into structured fields such as project type, scope, deliverables, staffing, and risks."],
            ["Why is user clarification still needed?", "AI may not know exact start date, team size, duration, or delivery model. Human confirmation improves accuracy."],
            ["Why SQLite?", "SQLite is simple for local demos and small deployments. For enterprise scale, PostgreSQL would be stronger."],
            ["What is the role of Flask?", "Flask serves the web pages and backend APIs."],
            ["How are documents processed?", "Files are validated, saved, routed to a parser, and converted into extracted text."],
            ["What does Claude return?", "Structured JSON with project type, client name, scope, deliverables, team requirements, confidence, and missing fields."],
            ["How are workbooks generated?", "The backend reads project data from SQLite and uses workbook generator services built with openpyxl."],
            ["Did the current package use Harness?", "The package is Harness-ready but does not include Harness execution evidence. Harness is best described as the planned or recommended CI/CD orchestration layer."],
            ["What would Harness add?", "Automated builds, tests, secret handling, approvals, deployments, health checks, audit trail, and rollback."],
        ],
        widths=[Inches(2.2), Inches(4.8)],
    )

    add_heading(doc, "23. Final Summary", 1)
    add_body(
        doc,
        "Project Aura is a complete project planning assistant that combines document extraction, AI analysis, project clarification, SQLite persistence, resource recommendation, workbook generation, and optional platform delivery. "
        "Its current local implementation is suitable for demos and controlled use. Its architecture can be extended toward enterprise deployment by adding persistent object storage, authentication, stronger database infrastructure, RAG/vector search, formal background jobs, and Harness-managed CI/CD."
    )

    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    create_document()
    print(f"Created {OUTPUT_FILE}")
