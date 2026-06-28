"""
DOCX document parser using python-docx.
Extracts text content and structure from DOCX files.
"""

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import logging
from typing import Dict, List, Any

logger = logging.getLogger(__name__)


class DOCXParser:
    """
    Parser for DOCX documents.
    Extracts text content, tables, and document structure.
    """

    @staticmethod
    def parse(file_path: str, max_length: int = 50000) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract its content.
        
        Args:
            file_path: Path to the DOCX file
            max_length: Maximum characters to extract (default: 50000)
        
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - tables: List of extracted tables
                - metadata: Document metadata
                - success: Boolean indicating parsing success
                - error: Error message if parsing failed
        """
        try:
            doc = Document(file_path)

            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or 'Unknown',
                'author': doc.core_properties.author or 'Unknown',
                'subject': doc.core_properties.subject or 'Unknown',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }

            # Extract text content with structure
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading levels if available
                    text_content.append(para.text)

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    'table_num': table_idx,
                    'data': table_data,
                    'rows': len(table.rows),
                    'cols': len(table.columns)
                })

            # Combine all extracted text
            combined_text = '\n\n'.join(text_content)
            
            # Add table summary to text
            if tables:
                combined_text += '\n\n--- TABLES ---\n'
                for t in tables:
                    combined_text += f"\nTable {t['table_num']} ({t['rows']}x{t['cols']}):\n"
                    for row in t['data'][:10]:  # Limit rows shown
                        combined_text += ' | '.join(row) + '\n'
                    if len(t['data']) > 10:
                        combined_text += f"... ({len(t['data']) - 10} more rows)\n"

            # Truncate if necessary
            if len(combined_text) > max_length:
                combined_text = combined_text[:max_length] + '\n\n[Content truncated...]'

            return {
                'success': True,
                'text': combined_text,
                'tables': tables,
                'metadata': metadata,
                'page_count': 'N/A',  # DOCX doesn't have fixed pages
                'has_tables': len(tables) > 0,
                'error': None
            }

        except FileNotFoundError:
            error_msg = f"DOCX file not found: {file_path}"
            logger.error(error_msg)
            return {
                'success': False,
                'text': '',
                'tables': [],
                'metadata': {},
                'page_count': 'N/A',
                'has_tables': False,
                'error': error_msg
            }

        except Exception as e:
            error_msg = f"Error parsing DOCX: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'text': '',
                'tables': [],
                'metadata': {},
                'page_count': 'N/A',
                'has_tables': False,
                'error': error_msg
            }
